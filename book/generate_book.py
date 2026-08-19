from pathlib import Path

import frontmatter
import yaml
from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]

PROGRAM = ROOT / "_data" / "program.yml"
TALKS_DIR = ROOT / "_talks"
POSTERS_DIR = ROOT / "_posters"

TALK_DOCX_DIR = ROOT / "source_abstracts"
POSTER_DOCX_DIR = ROOT / "source_posters"

OUTPUT = ROOT / "book-of-abstracts.docx"

def set_default_fonts(doc):
    styles = doc.styles

    for style_name in [
        "Normal",
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
    ]:
        if style_name not in [s.name for s in styles]:
            continue
        style = styles[style_name]
        style.font.name = "Calibri"

        # Important for Word compatibility
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri") 

def add_page_break(composer):
    p = composer.doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def page_break_document():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return doc

def load_items(directory):
    items = {}

    for path in directory.glob("*.md"):
        post = frontmatter.load(path)

        name = post.get("name")
        if not name:
            continue

        items[name] = {
            "path": path,
            **post.metadata,
        }

    return items


def add_heading_document(text, level=1):
    """Create a tiny DOCX containing just a heading."""
    doc = Document()
    doc.add_heading(text, level=level)
    return doc


def append_docx(composer, path, page_break_after=True):
    if not path.exists():
        print(f"WARNING: missing DOCX: {path}")
        return False

    composer.append(Document(path))

    composer.append(page_break_document())

    return True

def toc_document():
    doc = Document()

    heading = doc.add_heading("Contents", level=1)
    heading.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    add_toc(p)

    return doc

def add_toc(paragraph):
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "Right-click and update field"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)

def force_calibri(path):
    doc = Document(path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
                        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    doc.save(path)


def main():
    talks = load_items(TALKS_DIR)

    with PROGRAM.open(encoding="utf-8") as f:
        program = yaml.safe_load(f)

    FRONTMATTER_1 = ROOT / "book" / "Booklet-frontmatter-1.docx"
    FRONTMATTER_2 = ROOT / "book" / "Booklet-frontmatter-2.docx"
#    FRONTMATTER_3 = ROOT / "book" / "Booklet-frontmatter-3.docx"

    master = Document(FRONTMATTER_1)
    set_default_fonts(master)
    
    composer = Composer(master)
    
    composer.append(Document(FRONTMATTER_2))

    add_page_break(composer)
    composer.append(toc_document())

    # --------------------------------------------------
    # Oral contributions
    # --------------------------------------------------

    composer.append(add_heading_document("Oral Contributions", 1))

    previous_track = None

    for day in program["days"]:
        for room in day.get("rooms", []):
            for entry in room.get("talks", []):

                name = entry["name"]
                talk = talks.get(name)

                presenter = ", ".join(talk.get("speakers", []))
                
                heading_text = name
                if presenter:
                    heading_text += f" — {presenter}"
                # Programme items such as coffee, lunch, social events, etc.
                # do not necessarily have a _talks entry.
                if talk is None:
                    continue

                if talk.get("exclude_from_book") is True:
                    continue

                source = talk.get("abstract_source")
                if not source:
                    print(f"WARNING: no abstract_source for: {name}")
                    continue

                track = talk.get("track")

                # Add session heading whenever the track changes.
                if track and track != previous_track:
                    composer.append(
                        add_heading_document(track, 2)
                    )
                    previous_track = track

                path = TALK_DOCX_DIR / source

                print(f"TALK: {track}: {name}")
                composer.append(add_heading_document(heading_text, level=2))
                append_docx(composer, path)

    # --------------------------------------------------
    # Posters
    # --------------------------------------------------

    posters = []

    for path in POSTERS_DIR.glob("*.md"):
        post = frontmatter.load(path)

        posters.append({
            "number": int(post.get("number", 9999)),
            "name": post.get("name", path.stem),
            "abstract_source": post.get("abstract_source"),
        })

    posters.sort(key=lambda x: x["number"])
    previous_track = None

    add_page_break(composer)
    composer.append(
        add_heading_document("Poster Contributions", 1)
    )
    for poster in posters:
        source = poster["abstract_source"]

        if not source:
            print(
                f"WARNING: no abstract_source for poster "
                f"{poster['number']}: {poster['name']}"
            )
            continue
        title = f"Poster {poster['number']}: {poster['name']}"

        presenter = poster.get("presenter", "")
        
        heading_text = f"Poster {poster['number']}: {poster['name']}"
        
        if presenter:
            heading_text += f" — {presenter}"

        print(title)

        composer.append(add_heading_document(heading_text, level=2))
        append_docx(
            composer,
            POSTER_DOCX_DIR / source,
        )
    composer.save(OUTPUT)
    force_calibri(OUTPUT)
    print()
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
